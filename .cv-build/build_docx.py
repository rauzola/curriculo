# -*- coding: utf-8 -*-
"""Gera um curriculo .docx ATS-friendly a partir de um JSON (RESUME_SCHEMA)."""
import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

data = json.load(open(sys.argv[1], encoding="utf-8"))
out = sys.argv[2]

ACCENT = RGBColor(0x0B, 0x5F, 0x8A)  # azul profissional

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.5)
    s.bottom_margin = Inches(0.5)
    s.left_margin = Inches(0.7)
    s.right_margin = Inches(0.7)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)


def sp(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "2")
    b.set(qn("w:color"), "0B5F8A")
    pbdr.append(b)
    pPr.append(pbdr)


def section(title):
    p = doc.add_paragraph()
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    bottom_border(p)
    sp(p, before=8, after=4)


# Cabecalho
p = doc.add_paragraph()
r = p.add_run(data["nome"])
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = ACCENT
sp(p, after=0)

p = doc.add_paragraph()
r = p.add_run(data["tituloAlvo"])
r.bold = True
r.font.size = Pt(12)
sp(p, after=2)

c = data.get("contatos", {})
bits = [data.get("localizacao"), c.get("email"), c.get("telefone"),
        c.get("linkedin"), c.get("github"), c.get("portfolio"),
        c.get("curriculoOnline")]
bits = [b for b in bits if b]
p = doc.add_paragraph()
r = p.add_run("  •  ".join(bits))
r.font.size = Pt(9)
sp(p, after=6)

# Resumo
section("Resumo")
p = doc.add_paragraph(data["resumo"])
sp(p, after=4, line=1.1)

# Competencias
if data.get("competencias"):
    section("Competências")
    for cat in data["competencias"]:
        p = doc.add_paragraph()
        r = p.add_run(cat["categoria"] + ": ")
        r.bold = True
        p.add_run(", ".join(cat["itens"]))
        sp(p, after=2)

# Experiencia
section("Experiência Profissional")
for e in data.get("experiencias", []):
    p = doc.add_paragraph()
    r = p.add_run(e["cargo"])
    r.bold = True
    r.font.size = Pt(11)
    sp(p, before=5, after=0)
    sub = " | ".join([x for x in [e.get("empresa"), e.get("periodo"), e.get("local")] if x])
    p2 = doc.add_paragraph()
    r2 = p2.add_run(sub)
    r2.italic = True
    r2.font.size = Pt(9.5)
    sp(p2, after=2)
    for b in e.get("bullets", []):
        bp = doc.add_paragraph(b, style="List Bullet")
        sp(bp, after=1)

# Projetos
if data.get("projetos"):
    section("Projetos em Destaque")
    for pr in data["projetos"]:
        p = doc.add_paragraph()
        r = p.add_run(pr["nome"])
        r.bold = True
        if pr.get("url"):
            ru = p.add_run("  —  " + pr["url"])
            ru.font.size = Pt(9)
        sp(p, before=3, after=0)
        desc = pr["descricao"]
        p2 = doc.add_paragraph(desc)
        sp(p2, after=0)
        if pr.get("stack"):
            p3 = doc.add_paragraph()
            r3 = p3.add_run("Stack: " + pr["stack"])
            r3.italic = True
            r3.font.size = Pt(9)
            sp(p3, after=2)

# Formacao
if data.get("formacao"):
    section("Formação Acadêmica")
    for f in data["formacao"]:
        p = doc.add_paragraph()
        r = p.add_run(f["curso"])
        r.bold = True
        p.add_run("  —  %s · %s" % (f["instituicao"], f["periodo"]))
        sp(p, after=1)

# Cursos
if data.get("cursos"):
    section("Cursos e Certificações")
    for cu in data["cursos"]:
        bits = [cu.get("nome"), cu.get("instituicao"), cu.get("periodo")]
        p = doc.add_paragraph(" — ".join([b for b in bits if b]), style="List Bullet")
        sp(p, after=1)

# Idiomas
if data.get("idiomas"):
    section("Idiomas")
    p = doc.add_paragraph("; ".join("%s: %s" % (i["idioma"], i["nivel"]) for i in data["idiomas"]))
    sp(p, after=2)

doc.save(out)
print("DOCX OK ->", out)
